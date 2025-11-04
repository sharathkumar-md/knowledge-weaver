"""DOCX document processor"""

from pathlib import Path
from typing import Dict, Any
import docx
from .base_processor import BaseProcessor
from loguru import logger


class DOCXProcessor(BaseProcessor):
    """Processor for Microsoft Word documents (.docx)"""

    SUPPORTED_EXTENSIONS = {'.docx'}

    def can_process(self, file_path: Path) -> bool:
        """Check if file is a DOCX document"""
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)

            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            text = '\n\n'.join(text_parts)
            logger.info(f"Extracted DOCX: {file_path.name}")

            return text

        except Exception as e:
            raise RuntimeError(f"Failed to process DOCX file {file_path}: {e}")

    def _create_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Create metadata with DOCX-specific info"""
        metadata = super()._create_metadata(file_path)

        try:
            doc = docx.Document(file_path)

            # Extract document properties
            core_props = doc.core_properties

            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.keywords:
                metadata['keywords'] = core_props.keywords
            if core_props.created:
                metadata['doc_created'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['doc_modified'] = core_props.modified.isoformat()

            # Count paragraphs and tables
            metadata['num_paragraphs'] = len(doc.paragraphs)
            metadata['num_tables'] = len(doc.tables)

        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata from {file_path.name}: {e}")

        return metadata
